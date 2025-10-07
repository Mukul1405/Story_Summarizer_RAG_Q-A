"""
Streamlit app: Azure DevOps Story Summarizer + Conversational QA

Save as: app.py
Run:    streamlit run app.py

Pip install recommended:
pip install streamlit requests python-docx PyPDF2 langchain sentence-transformers faiss-cpu transformers accelerate
# If using langchain_groq / ChatGroq:
pip install langchain-groq   # if available (package name may differ). See note below.

ENV:
- Optionally set:
  AZURE_DEVOPS_ORG: your_org
  AZURE_DEVOPS_PROJECT: your_project
  AZURE_DEVOPS_PAT: personal_access_token
  GROQ_API_KEY: your_groq_api_key (if using Groq)

NOTE:
- This script tries to use ChatGroq from langchain_groq (if installed).
  If not installed/available, you'll see a clear message and can switch to a HuggingFace LLM.
- Embeddings: Uses HuggingFace 'sentence-transformers/all-MiniLM-L6-v2' via langchain's HuggingFaceEmbeddings.
- Vectorstore: FAISS (CPU) stored in memory (not persisted). You can extend to persist disk if needed.
"""

import os
import json
import base64
import requests
from typing import List, Tuple
from dotenv import load_dotenv

load_dotenv()

import streamlit as st

# Text extraction
from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile

# docx/pdf parsing
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

# LangChain pieces
from langchain.text_splitter import CharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain, RetrievalQA
from langchain.memory import ConversationBufferMemory
from langchain.docstore.document import Document

# optional: ChatGroq (if available in your environment)
USE_GROQ = False
try:
    # attempt import - may fail if package not installed
    from langchain_groq import ChatGroq  # type: ignore
    USE_GROQ = True
except Exception:
    USE_GROQ = False

# -----------------------------
# Helper functions
# -----------------------------

def azure_fetch_work_item(org: str, project: str, work_item_id: str, pat: str) -> Tuple[str,str]:
    """
    Fetches a work item from Azure DevOps and returns (description_text, acceptance_criteria_text)
    - work_item_id can be numeric or string
    - Uses Basic auth with PAT (username blank)
    """
    if not all([org, project, work_item_id, pat]):
        raise ValueError("Azure DevOps org, project, work_item_id and PAT are required.")

    url = f"https://dev.azure.com/{org}/{project}/_apis/wit/workitems/{work_item_id}?api-version=7.1"
    resp = requests.get(url, auth=('', pat))
    if resp.status_code != 200:
        raise RuntimeError(f"Failed to fetch work item: {resp.status_code} - {resp.text}")
    data = resp.json()

    fields = data.get("fields", {})
    # Common fields: System.Description, Microsoft.VSTS.Common.AcceptanceCriteria (or custom)
    description = fields.get("System.Description") or fields.get("Microsoft.VSTS.TCM.ReproSteps") or ""
    acceptance = fields.get("Microsoft.VSTS.Common.AcceptanceCriteria") or fields.get("Custom.AcceptanceCriteria") or ""

    # Some descriptions come as HTML. Quick strip of tags (rudimentary).
    def strip_html(html: str) -> str:
        if not html:
            return ""
        # rudimentary strip (keeps text content)
        try:
            # a light removal of HTML tags
            from html import unescape
            import re
            text = re.sub(r'<[^>]+>', '', html)
            return unescape(text).strip()
        except Exception:
            return html

    return strip_html(description), strip_html(acceptance)


def extract_text_from_file(uploaded_file) -> str:
    """
    Accepts a Streamlit UploadedFile and extracts plain text.
    Supported: .txt, .md, .docx, .pdf
    """
    name = uploaded_file.name.lower()
    content = uploaded_file.read()
    if name.endswith(".txt") or name.endswith(".md"):
        try:
            return content.decode('utf-8', errors='ignore')
        except Exception:
            return str(content)
    elif name.endswith(".docx"):
        if DocxDocument is None:
            raise RuntimeError("python-docx not installed. Install with `pip install python-docx`")
        # need to load into a temp file path Document expects a filename or file-like
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.write(content)
        tmp.flush()
        tmp.close()
        doc = DocxDocument(tmp.name)
        text = "\n".join([p.text for p in doc.paragraphs])
        os.unlink(tmp.name)
        return text
    elif name.endswith(".pdf"):
        if PyPDF2 is None:
            raise RuntimeError("PyPDF2 not installed. Install with `pip install PyPDF2`")
        tmp = BytesIO(content)
        reader = PyPDF2.PdfReader(tmp)
        pages = []
        for p in reader.pages:
            try:
                pages.append(p.extract_text() or "")
            except Exception:
                pages.append("")
        return "\n".join(pages)
    else:
        # Try to decode anyway
        try:
            return content.decode('utf-8', errors='ignore')
        except Exception:
            return str(content)


def build_docs_from_texts(title_text_pairs: List[Tuple[str, str]]) -> List[Document]:
    """
    Convert [(title, text)] to LangChain Document objects with metadata.
    """
    docs = []
    for title, text in title_text_pairs:
        if not text:
            continue
        # simple text splitter to create smaller docs
        splitter = CharacterTextSplitter(chunk_size=1200, chunk_overlap=150)
        chunks = splitter.split_text(text)
        for i, chunk in enumerate(chunks):
            docs.append(Document(page_content=chunk, metadata={"source_title": title, "chunk": i}))
    return docs


def get_embedding_model(model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
    """
    Returns a HuggingFaceEmbeddings instance (langchain).
    """
    # HuggingFaceEmbeddings from langchain will try to load the model locally or via huggingface.
    return HuggingFaceEmbeddings(model_name=model_name)


def create_vectorstore_from_docs(docs: List[Document], embeddings) -> FAISS:
    """
    Build a FAISS vectorstore from the given langchain Documents and embeddings.
    """
    if not docs:
        raise ValueError("No documents supplied to embed.")
    vectorstore = FAISS.from_documents(docs, embedding=embeddings)
    return vectorstore


def get_llm():
    """
    Returns an LLM instance. Prefers ChatGroq (if available and GROQ_API_KEY exists).
    If not available, raises informative error — user may swap in an available LLM.
    """
    if os.getenv("GROQ_API_KEY"):
        # Construct ChatGroq as per langchain_groq API
        api_key = os.getenv("GROQ_API_KEY")
        # default model name — change as needed
        llm = ChatGroq(api_key=api_key, model=os.getenv("model"))  # placeholder model name; update per provider docs
        return llm
    else:
        # Helpful error explaining how to enable
        raise RuntimeError(
            "Groq LLM (ChatGroq) is not available in this environment or GROQ_API_KEY not set. "
            "To use Groq: install langchain_groq (if available) and set environment variable GROQ_API_KEY. "
            "Alternatively, modify get_llm() to return a HuggingFace or other LLM instance supported by your environment."
        )

# -----------------------------
# Streamlit UI & App Logic
# -----------------------------

st.set_page_config(page_title="Azure DevOps Story Summarizer (LangChain + Groq + HF Embeddings + FAISS)",
                   layout="wide")

st.title("Azure DevOps Story Summarizer — Ask questions about the story (LangChain + Groq + HF Embeddings + FAISS)")

st.markdown("""
**How it works (summary):**
1. Enter an Azure DevOps Story ID OR upload a document (.txt/.md/.docx/.pdf).  
2. The app fetches/extracts the description + acceptance criteria (or document text), embeds the content using HuggingFace embeddings, stores it in FAISS, and builds a conversational retrieval chain using an LLM (Groq expected).  
3. After embedding, you'll see a message telling you you can ask questions. Ask scenario/doubt questions — chat history will be preserved.
""")

# ✅ Prefill sidebar defaults BEFORE rendering widgets
if "org" not in st.session_state:
    st.session_state.org = os.getenv("AZURE_DEVOPS_ORG", "")
if "project" not in st.session_state:
    st.session_state.project = os.getenv("AZURE_DEVOPS_PROJECT", "")
if "pat" not in st.session_state:
    st.session_state.pat = os.getenv("AZURE_DEVOPS_PAT", "")
org_input = st.sidebar.text_input("Azure DevOps Organization (org)", key="org")
project_input = st.sidebar.text_input("Azure DevOps Project", key="project")
pat_input = st.sidebar.text_input("Azure DevOps PAT (personal access token)", key="pat", type="password")


grok_key_hint = "GROQ_API_KEY present in environment: " + ("Yes" if os.getenv("GROQ_API_KEY") else "No")
st.sidebar.write(grok_key_hint)

# Main inputs
col1, col2 = st.columns([2,1])
with col1:
    story_id = st.text_input("Enter Azure DevOps Story/Work Item ID (leave empty if uploading document):")
    uploaded_file = st.file_uploader("Or upload a document (.txt, .md, .docx, .pdf)", type=["txt","md","docx","pdf"])
    process_btn = st.button("Fetch & Embed / Process")

with col2:
    st.markdown("**Status**")
    status_placeholder = st.empty()
    st.markdown("---")
    st.markdown("**Chat**")
    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []  # list of (user, assistant)
    if "vectorstore" not in st.session_state:
        st.session_state["vectorstore"] = None
    if "retrieval_chain" not in st.session_state:
        st.session_state["retrieval_chain"] = None
    if "docs_sources" not in st.session_state:
        st.session_state["docs_sources"] = []

# Processing step: fetch or extract & embed
if process_btn:
    try:
        status_placeholder.info("Processing... extracting text and building embeddings. This may take a moment.")
        title_text_pairs = []
        if story_id:
            # fetch from Azure DevOps
            description, acceptance = azure_fetch_work_item(org_input, project_input, story_id, pat_input)
            title_text_pairs.append((f"WorkItem-{story_id}-Description", description))
            if acceptance:
                title_text_pairs.append((f"WorkItem-{story_id}-AcceptanceCriteria", acceptance))
            if not description and not acceptance:
                st.warning("Fetched work item but description and acceptance criteria are empty.")
        if uploaded_file is not None:
            doc_text = extract_text_from_file(uploaded_file)
            title_text_pairs.append((uploaded_file.name, doc_text))

        if not title_text_pairs:
            st.error("No inputs found. Provide a story id or upload a document.")
        else:
            # Build docs
            docs = build_docs_from_texts(title_text_pairs)
            if not docs:
                st.error("No content extracted to embed.")
            else:
                status_placeholder.info("Creating embeddings (HuggingFace) and FAISS index...")
                embedder = get_embedding_model()
                vs = create_vectorstore_from_docs(docs, embedder)
                st.session_state["vectorstore"] = vs
                st.session_state["docs_sources"] = [t for t,_ in title_text_pairs]

                # Prepare conversational retrieval chain with LLM
                try:
                    llm = get_llm()
                    # Conversation memory
                    memory = ConversationBufferMemory(return_messages=True, memory_key="chat_history")
                    retriever = vs.as_retriever(search_type="similarity", search_kwargs={"k":4})
                    # ConversationalRetrievalChain expects an LLM with .generate or .__call__
                    conv_chain = ConversationalRetrievalChain.from_llm(llm=llm, retriever=retriever, memory=memory)
                    st.session_state["retrieval_chain"] = conv_chain
                    status_placeholder.success("Embedding complete. You can now ask questions from the embedded document/story.")
                except Exception as e:
                    # LLM setup failed — keep vectorstore but inform how to proceed
                    st.session_state["retrieval_chain"] = None
                    status_placeholder.warning(f"Embeddings created and stored in FAISS but failed to initialise Groq LLM: {str(e)}")
                    st.info("To enable Groq LLM: install the Groq provider integration (if available) and set GROQ_API_KEY env var, or modify get_llm() to use another LLM (HuggingFace local model).")
    except Exception as e:
        st.exception(e)

# Chat UI
st.markdown("---")
st.subheader("Ask questions about the embedded story / document")

if st.session_state.get("vectorstore") is None:
    st.info("No content embedded yet. Enter a story ID or upload a document and press 'Fetch & Embed / Process'.")
else:
    if st.session_state.get("retrieval_chain") is None:
        st.warning("Content embedded into FAISS, but LLM retrieval chain not ready. See sidebar for Groq API key or check logs.")
        st.write("You can still embed content, but to ask questions the app needs a working LLM. Modify get_llm() to return an LLM that's available in your environment.")
    else:
        user_question = st.text_input("Ask a question (e.g., 'Summarize in layman terms', 'List important scenarios', 'What are edge cases?')", key="user_q_input")
        if st.button("Send Question"):
            if not user_question:
                st.warning("Please enter a question.")
            else:
                try:
                    chain: ConversationalRetrievalChain = st.session_state["retrieval_chain"]
                    # Run conversational chain
                    result = chain({"question": user_question})
                    answer = result.get("answer") or result.get("output_text") or str(result)
                    # Append to session history
                    st.session_state.chat_history.append(("user", user_question))
                    st.session_state.chat_history.append(("assistant", answer))
                except Exception as e:
                    st.exception(f"Error when running retrieval chain: {e}")

        # Display chat history
        if st.session_state.chat_history:
            st.markdown("**Chat history**")
            for role, text in st.session_state.chat_history[::-1]:  # reverse show recent first
                if role == "user":
                    st.markdown(f"**You:** {text}")
                else:
                    st.markdown(f"**Assistant:** {text}")

# Small utilities: show embedded docs metadata and a "clear" button
st.markdown("---")
st.subheader("Debug / Utilities")
if st.button("Show embedded document sources"):
    sources = st.session_state.get("docs_sources", [])
    if sources:
        st.json(sources)
    else:
        st.write("No embedded sources stored in session.")

if st.button("Reset session (clear vectorstore and chat)"):
    st.session_state.pop("vectorstore", None)
    st.session_state.pop("retrieval_chain", None)
    st.session_state.pop("chat_history", None)
    st.session_state.pop("docs_sources", None)
    st.success("Session cleared. Reload page to ensure clean state.")

st.markdown("""
---
**Developer / Notes**

- The script expects `sentence-transformers/all-MiniLM-L6-v2` (or similar) for embeddings. If you want another embedding model, change the `get_embedding_model()` call.
- Groq LLM usage depends on `langchain_groq` and a valid `GROQ_API_KEY`. If you don't have Groq enabled, replace `get_llm()` with a local HuggingFace LLM using `transformers` or use another LLM provider supported by LangChain.
- FAISS index is in-memory. To persist, you can call `faiss.write_index` to disk and reload later; extend the code as needed.
""")
