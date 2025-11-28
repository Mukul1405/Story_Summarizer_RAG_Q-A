import os
import re
import json
import requests
from typing import List, Tuple
import time
import base64

import streamlit as st
from dotenv import load_dotenv

from io import BytesIO
from tempfile import NamedTemporaryFile

# Jira integration helpers (local)
try:
    from jira_integration import jira_fetch_issue, jira_fetch_comments, build_jira_issue_key, jira_fetch_linked_prs
except Exception:
    jira_fetch_issue = None
    jira_fetch_comments = None
    build_jira_issue_key = None
# Optionally available libs for docx/pdf/html parsing
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    from bs4 import BeautifulSoup
except Exception:
    BeautifulSoup = None

# Excel handling
try:
    import pandas as pd
    import openpyxl
except Exception:
    pd = None
    openpyxl = None

# LangChain pieces (standard names)
try:
    from langchain_text_splitters import CharacterTextSplitter
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS
    from langchain_classic.chains import ConversationalRetrievalChain
    from langchain_classic.memory import ConversationBufferMemory
    from langchain_core.documents import Document
except Exception as e:
    st.error("Required langchain packages not found. Please install langchain and related packages.")
    raise

# Try to import Groq provider (optional)
USE_GROQ = False
ChatGroq = None
try:
    from langchain_groq import ChatGroq  # type: ignore
    USE_GROQ = True
except Exception:
    USE_GROQ = False
    ChatGroq = None

# Groq client for test case generation
try:
    from groq import Groq
    GROQ_CLIENT_AVAILABLE = True
except Exception:
    GROQ_CLIENT_AVAILABLE = False
    Groq = None

# -------------------------
# Utility functions
# -------------------------

def load_env_to_session():
    load_dotenv()
    if "jira_base_url" not in st.session_state:
        # Jira base URL can be a full browse URL or the base host
        st.session_state["jira_base_url"] = os.getenv("JIRA_BASE_URL", "")
    if "jira_pat" not in st.session_state:
        st.session_state["jira_pat"] = os.getenv("JIRA_PAT", "")
    if "jira_email" not in st.session_state:
        st.session_state["jira_email"] = os.getenv("JIRA_EMAIL", "")
    if "groq_key" not in st.session_state:
        st.session_state["groq_key"] = os.getenv("GROQ_API_KEY", "")
    if "vectorstore" not in st.session_state:
        st.session_state["vectorstore"] = None
    if "retrieval_chain" not in st.session_state:
        st.session_state["retrieval_chain"] = None
    if "docs_sources" not in st.session_state:
        st.session_state["docs_sources"] = []
    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []
    if "qa_history" not in st.session_state:
        st.session_state["qa_history"] = []
    if "flow_diagram" not in st.session_state:
        st.session_state["flow_diagram"] = None
    if "story_data" not in st.session_state:
        st.session_state["story_data"] = []
    if "test_cases_excel" not in st.session_state:
        st.session_state["test_cases_excel"] = None

def clean_org_for_api(org: str) -> str:
    if not org:
        return ""
    org = org.strip()
    m = re.search(r"dev\.azure\.com/([^/]+)", org, flags=re.IGNORECASE)
    if m:
        return m.group(1)
    m = re.search(r"https?://([^./]+)\.visualstudio\.com", org, flags=re.IGNORECASE)
    if m:
        return m.group(1)
    if org.startswith("http"):
        parts = [p for p in org.split("/") if p]
        if parts:
            return parts[-1]
    return org

def strip_html_to_text(html: str) -> str:
    if not html:
        return ""
    if BeautifulSoup:
        try:
            soup = BeautifulSoup(html, "html.parser")
            for s in soup(["script", "style"]):
                s.decompose()
            text = soup.get_text(separator="\n")
            text = re.sub(r"\n\s*\n+", "\n\n", text).strip()
            return text
        except Exception:
            pass
    text = re.sub(r"<[^>]+>", "", html)
    return text.strip()

def azure_fetch_work_item(org: str, project: str, work_item_id: str, pat: str) -> Tuple[str, str]:
    if not all([org, project, work_item_id, pat]):
        raise ValueError("Azure DevOps org, project, work_item_id and PAT are required.")
    org_clean = clean_org_for_api(org)
    base_url = f"https://dev.azure.com/{org_clean}"
    url = f"{base_url}/{project}/_apis/wit/workitems/{work_item_id}?api-version=7.1"
    resp = requests.get(url, auth=("", pat))
    if resp.status_code != 200:
        ct = resp.headers.get("content-type", "")
        if resp.status_code in (401, 403) or ("text/html" in ct and "Sign In" in resp.text[:200]):
            raise RuntimeError(
                f"Authentication failed when fetching work item {work_item_id}. "
                "Check your PAT (permissions: Work Items read) and that org/project are correct."
            )
        raise RuntimeError(f"Failed to fetch work item {work_item_id}: {resp.status_code} - {resp.text[:300]}")
    data = resp.json()
    fields = data.get("fields", {})
    description = fields.get("System.Description") or fields.get("Microsoft.VSTS.TCM.ReproSteps") or ""
    acceptance = fields.get("Microsoft.VSTS.Common.AcceptanceCriteria") or fields.get("Custom.AcceptanceCriteria") or ""
    return strip_html_to_text(description), strip_html_to_text(acceptance)

def azure_fetch_work_item_comments(org: str, project: str, work_item_id: str, pat: str) -> List[dict]:
    """
    Fetch ALL comments for a work item with pagination and priority ranking.
    Returns list of dicts: [{'author': str, 'text': str, 'created': str, 'weight': int}]
    """
    try:
        org_clean = clean_org_for_api(org)
        base_url = f"https://dev.azure.com/{org_clean}"
        url = f"{base_url}/{project}/_apis/wit/workItems/{work_item_id}/comments?api-version=7.1-preview.3"

        all_comments = []
        continuation_token = None

        while True:
            headers = {}
            if continuation_token:
                headers["x-ms-continuationtoken"] = continuation_token

            resp = requests.get(url, auth=("", pat), headers=headers)
            if resp.status_code == 404:
                break  # No comments
            resp.raise_for_status()

            data = resp.json()
            comments = data.get("comments", [])
            for c in comments:
                author = c.get("createdBy", {}).get("displayName", "Unknown")
                text = strip_html_to_text(c.get("text", "")).strip()
                created = c.get("createdDate", "")
                if text:
                    weight = 1
                    name_lower = author.lower()
                    if any(k in name_lower for k in ["dev", "developer", "engineer"]):
                        weight = 3
                    elif any(k in name_lower for k in ["qa", "test"]):
                        weight = 2
                    elif any(k in name_lower for k in ["po", "manager", "lead"]):
                        weight = 1

                    all_comments.append({
                        "author": author,
                        "text": text,
                        "created": created,
                        "weight": weight
                    })

            continuation_token = resp.headers.get("x-ms-continuationtoken")
            if not continuation_token:
                break  # no more pages

        # Sort by weight (desc) then recency
        all_comments.sort(key=lambda x: (x["weight"], x["created"]), reverse=True)
        return all_comments

    except Exception as e:
        st.warning(f"Comment fetch failed for {work_item_id}: {e}")
        return []

import urllib.parse

def azure_fetch_linked_prs(org: str, project: str, work_item_id: str, pat: str) -> List[dict]:
    """
    Fetch Pull Requests linked to a given Azure DevOps work item.
    Supports 3-part vstfs PullRequestId format:
    vstfs:///Git/PullRequestId/<projectId>%2F<repoId>%2F<prId>
    Returns: list of dicts [{ 'projectId': str, 'repoId': str, 'prId': str }]
    """
    try:
        org_clean = clean_org_for_api(org)
        base_url = f"https://dev.azure.com/{org_clean}"
        url = f"{base_url}/{project}/_apis/wit/workItems/{work_item_id}?$expand=relations&api-version=7.1"

        resp = requests.get(url, auth=("", pat))
        resp.raise_for_status()
        data = resp.json()

        relations = data.get("relations", [])
        prs = []

        for rel in relations:
            url_val = rel.get("url", "")
            if "PullRequestId" in url_val:
                encoded = url_val.split("PullRequestId/")[-1]
                decoded = urllib.parse.unquote(encoded)  # decode %2F → /
                parts = decoded.split("/")
                if len(parts) == 3:  # ProjectID / RepoID / PRID
                    project_id, repo_id, pr_id = parts
                    prs.append({
                        "projectId": project_id,
                        "repoId": repo_id,
                        "prId": pr_id
                    })
        return prs
    except Exception as e:
        st.warning(f"PR fetch failed for {work_item_id}: {e}")
        return []

def azure_fetch_pr_changes(org: str, project_id: str, repo_id: str, pr_id: str, pat: str) -> List[str]:
    """
    Fetch changed files for a PR (latest iteration)
    Supports both 'changes' and 'changeEntries' keys
    """
    try:
        org_clean = clean_org_for_api(org)
        base_url = f"https://dev.azure.com/{org_clean}"

        # Get latest iteration
        iter_url = f"{base_url}/{project_id}/_apis/git/repositories/{repo_id}/pullRequests/{pr_id}/iterations?api-version=7.1"
        iter_resp = requests.get(iter_url, auth=("", pat))
        iter_resp.raise_for_status()
        iter_data = iter_resp.json()
        iterations = iter_data.get("value", [])
        latest_iter = iterations[-1]["id"] if iterations else 1

        # Get changes for that iteration
        change_url = f"{base_url}/{project_id}/_apis/git/repositories/{repo_id}/pullRequests/{pr_id}/iterations/{latest_iter}/changes?api-version=7.1"
        resp = requests.get(change_url, auth=("", pat))
        resp.raise_for_status()
        data = resp.json()

        # ✅ handle both response formats
        entries = data.get("changes") or data.get("changeEntries") or []
        result = []
        for c in entries:
            item = c.get("item", {})
            path = item.get("path", "")
            change_type = c.get("changeType", "edit")
            if path:
                result.append(f"{change_type.title()}: {path}")

        return result

    except Exception as e:
        st.warning(f"Failed to fetch PR changes: {e}")
        return []

def generate_pr_test_focus(org: str, project: str, story_id: str, pat: str, story_text: str) -> str:
    """
    Analyze linked PRs + code changes and generate Smart Testing Focus report using LLM.
    """

    
    try:
        prs = azure_fetch_linked_prs(org, project, story_id, pat)
        if not prs:
            return f"*(No PRs linked to story {story_id})*"
        
        all_changes = []
        for pr in prs:
            repo_id = pr["repoId"]
            pr_id = pr["prId"]
            pr_changes = azure_fetch_pr_changes(org, pr["projectId"], pr["repoId"], pr["prId"], pat)
            all_changes.extend(pr_changes)
        
        if not all_changes:
            return "*(No changed files detected for linked PRs)*"
        
        # Limit change list for prompt brevity
        changes_text = "\n".join(all_changes[:50])
        
        prompt = f"""
You are an expert SDET and test architect.
Based on the following STORY CONTEXT and CODE CHANGES from linked PRs,
analyze where the system is most likely to break or any condition/check is missing and what areas need deep testing.

STORY CONTEXT:
{story_text[:1500]}

CODE CHANGES:
{changes_text}

Generate your analysis in this format:
### 🧠 Smart Testing Focus

**High Risk Areas**
- <short description>

**Potential Break Points**
- <short description>

**Suggested Smart Test Cases**
- <test case 1>
- <test case 2>

**Regression Scenarios**
- <scenario 1>
"""
        llm = get_llm_instance()
        raw = llm.invoke(prompt)
        return extract_llm_text(raw)
    except Exception as e:
        return f"Error during PR analysis: {e}"

def azure_run_wiql(org: str, project: str, wiql_query: str, pat: str) -> List[str]:
    if not all([org, project, wiql_query, pat]):
        raise ValueError("WIQL: org, project, query and PAT are required.")
    org_clean = clean_org_for_api(org)
    base_url = f"https://dev.azure.com/{org_clean}"
    url = f"{base_url}/{project}/_apis/wit/wiql?api-version=7.1"
    headers = {"Content-Type": "application/json"}
    body = {"query": wiql_query}
    resp = requests.post(url, auth=("", pat), headers=headers, data=json.dumps(body))
    if resp.status_code != 200:
        raise RuntimeError(f"WIQL query failed: {resp.status_code} - {resp.text[:500]}")
    data = resp.json()
    items = data.get("workItems", [])
    ids = [str(i.get("id")) for i in items if i.get("id")]
    return ids

def jira_run_jql(base_url: str, jql_query: str, pat: str, email: str = None) -> List[str]:
    """Run a JQL query against Jira Cloud and return a list of issue keys (e.g., BAW-123).

    This is a minimal helper; it returns up to 100 issue keys by default.
    """
    if not base_url or not jql_query or not pat:
        raise ValueError("Jira JQL: base_url, jql_query and JIRA PAT are required.")
    base = base_url.strip()
    if base.endswith("/"):
        base = base[:-1]
    url = f"{base}/rest/api/3/search"
    headers = {"Accept": "application/json", "Content-Type": "application/json"}
    auth = None
    if email or os.getenv("JIRA_EMAIL"):
        user = email or os.getenv("JIRA_EMAIL")
        auth = (user, pat)
    else:
        headers["Authorization"] = f"Bearer {pat}"

    body = {"jql": jql_query, "maxResults": 100, "fields": []}
    resp = requests.post(url, headers=headers, auth=auth, data=json.dumps(body))
    if resp.status_code != 200:
        raise RuntimeError(f"JQL query failed: {resp.status_code} - {resp.text[:300]}")
    data = resp.json()
    issues = data.get("issues", [])
    keys = [iss.get("key") for iss in issues if iss.get("key")]
    return keys

def extract_text_from_file(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    content = uploaded_file.read()
    if name.endswith((".txt", ".md")):
        try:
            return content.decode("utf-8", errors="ignore")
        except Exception:
            return str(content)
    if name.endswith(".docx"):
        if DocxDocument is None:
            raise RuntimeError("python-docx not installed. Install: pip install python-docx")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.write(content)
        tmp.flush()
        tmp.close()
        doc = DocxDocument(tmp.name)
        text = "\n".join([p.text for p in doc.paragraphs])
        try:
            os.unlink(tmp.name)
        except Exception:
            pass
        return text
    if name.endswith(".pdf"):
        if PyPDF2 is None:
            raise RuntimeError("PyPDF2 not installed. Install: pip install PyPDF2")
        tmp = BytesIO(content)
        reader = PyPDF2.PdfReader(tmp)
        pages = []
        for p in reader.pages:
            try:
                t = p.extract_text() or ""
            except Exception:
                t = ""
            pages.append(t)
        return "\n".join(pages)
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return str(content)

def build_docs_from_texts(title_text_pairs: List[Tuple[str, str]]) -> List[Document]:
    docs = []
    for title, text in title_text_pairs:
        if not text or not text.strip():
            continue
        content = f"Source: {title}\n\n{text.strip()}"
        docs.append(Document(page_content=content, metadata={"source": title}))
    return docs

def extract_llm_text(llm_result) -> str:
    if isinstance(llm_result, str):
        return llm_result
    
    # Try to get content attribute directly (for ChatGroq responses)
    try:
        if hasattr(llm_result, 'content'):
            return llm_result.content
    except Exception:
        pass
    
    try:
        gens = getattr(llm_result, "generations", None)
        if gens:
            g0 = gens[0][0]
            txt = getattr(g0, "text", None) or getattr(g0, "generation", None)
            if txt:
                return txt
    except Exception:
        pass
    if isinstance(llm_result, dict):
        for k in ("answer", "text", "content", "output_text"):
            if k in llm_result and llm_result[k]:
                return llm_result[k] if isinstance(llm_result[k], str) else str(llm_result[k])
        if "choices" in llm_result and isinstance(llm_result["choices"], (list, tuple)) and len(llm_result["choices"]) > 0:
            c = llm_result["choices"][0]
            if isinstance(c, dict):
                m = c.get("message") or c.get("text")
                if isinstance(m, dict):
                    return m.get("content") or str(m)
                if isinstance(m, str):
                    return m
    return str(llm_result)

def get_llm_instance():
    groq_key = st.session_state.get("groq_key") or os.getenv("GROQ_API_KEY", "")
    if groq_key:
        try:
            try:
                llm = ChatGroq(groq_api_key=groq_key, model_name=os.getenv("model", "llama-3.3-70b-versatile"), temperature=0.0)
            except TypeError:
                llm = ChatGroq(api_key=groq_key, model_name=os.getenv("model", "llama-3.3-70b-versatile"), temperature=0.0)
            return llm
        except Exception as e:
            raise RuntimeError(f"Failed to initialize ChatGroq: {e}")
    raise RuntimeError(
        "Groq LLM not configured or langchain_groq not installed. "
        "Set GROQ_API_KEY in .env and install langchain-groq, or modify get_llm_instance() to return a HuggingFace LLM."
    )

def strict_split(text, chunk_size=1000, chunk_overlap=200):
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end]
        if len(chunk) > chunk_size:
            chunk = chunk[:chunk_size]
        chunks.append(chunk)
        start += chunk_size - chunk_overlap
    return chunks

def generate_flow_diagram(title_text_pairs: List[Tuple[str, str]]) -> str:
    """
    Generate a Mermaid flow diagram using LLM to analyze story relationships
    """
    try:
        llm = get_llm_instance()
        
        # Prepare content for LLM analysis - limit content length
        stories_text = "\n\n".join([
            f"Story {idx+1}: {title}\nContent: {text[:1000]}"
            for idx, (title, text) in enumerate(title_text_pairs[:10])  # Limit to 10 stories
        ])
        
        prompt = f"""Analyze the following user stories and create a Mermaid flowchart diagram showing their relationships and workflow.

CRITICAL RULES:
1. Use ONLY these node types:
   - Rectangle: A[Text Here]
   - Rounded: A([Text Here])
   - Diamond: A{{{{Text Here}}}}  (use DOUBLE curly braces)
2. Keep node text SHORT (max 6 words)
3. Use SIMPLE arrow syntax: A --> B
4. Do NOT use special characters: ? : ; # & " in node text
5. Start with: graph LR
6. Output ONLY the Mermaid code, no explanations

STORIES:
{stories_text}

Generate Mermaid flowchart (graph LR format):"""

        result = llm.invoke(prompt)
        mermaid_code = extract_llm_text(result)
        
        # Clean up the response
        mermaid_code = mermaid_code.strip()
        mermaid_code = re.sub(r'^```mermaid\s*', '', mermaid_code)
        mermaid_code = re.sub(r'^```\s*', '', mermaid_code)
        mermaid_code = re.sub(r'```\s*$', '', mermaid_code)
        mermaid_code = mermaid_code.strip()
        
        # Ensure it starts with graph
        if not mermaid_code.startswith('graph'):
            mermaid_code = 'graph LR\n' + mermaid_code
        
        return mermaid_code
    
    except Exception as e:
        # Fallback: Create a simple diagram if LLM fails
        st.warning(f"Using fallback diagram (LLM error: {str(e)[:100]})")
        
        mermaid_code = "graph LR\n"
        mermaid_code += "    Start([Start])\n"
        
        # Limit to 8 stories
        for idx, (title, _) in enumerate(title_text_pairs[:8]):
            clean_title = title.replace('-', ' ').replace('_', ' ')[:40]
            clean_title = re.sub(r'[^\w\s]', '', clean_title)
            clean_title = ' '.join(clean_title.split())
            
            node_id = f"S{idx+1}"
            mermaid_code += f"    {node_id}[\"{clean_title}\"]\n"
            
            if idx == 0:
                mermaid_code += f"    Start --> {node_id}\n"
            else:
                mermaid_code += f"    S{idx} --> {node_id}\n"
        
        last_node = f"S{min(len(title_text_pairs), 8)}"
        mermaid_code += f"    End([Complete])\n"
        mermaid_code += f"    {last_node} --> End\n"
        
        return mermaid_code

# -------------------------
# TEST CASE GENERATION FUNCTIONS
# -------------------------

def sanitize_groq_output(text: str) -> str:
    """
    Fix common invalid JSON issues from AI output:
    - Escape unescaped backslashes
    - Replace invalid single quotes in keys with double quotes
    - Remove control chars if needed
    """
    # Escape unescaped backslashes (excluding already escaped ones)
    text = re.sub(r'(?<!\\)\\(?!["\\/bfnrtu])', r'\\\\', text)
    
    # Replace single quoted property names with double quotes
    text = re.sub(r"(?<=\{|\s)'([^']+)':", r'"\1":', text)

    # Remove control characters that can break JSON parsing
    text = re.sub(r"[\x00-\x1f\x7f]", "", text)

    return text

def make_groq_test_prompt(title: str, description: str, acceptance_criteria: str) -> str:
    prompt = f"""
You are a senior test-design engineer. For the user story below, produce a highly comprehensive, structured list of test cases.
Your output MUST be valid JSON (an array of objects). Each object should have:
  - test_case_id
  - priority (High/Medium/Low)
  - test_case (short title)
  - precondition
  - test_steps (ordered list of steps as an array)
  - expected_result

User Story Title:
{title}

Description:
{description}

Acceptance Criteria:
{acceptance_criteria}

RULES:
1. Generate complete set of test cases ranging(30-100 cases), covering all realistic, edge, boundary, integration, negative, performance, usability, security, accessibility, and error handling scenarios based on the story and typical project risks.
2. Go beyond acceptance criteria: include positive, negative, boundary value, data validation, regression, integration, compatibility, performance, security, usability, and accessibility test cases.
3. Assign priority based on risk, business value, and acceptance criteria.
4. Each expected_result must be 1–2 sentences, concise yet clear.
5. Ensure deep coverage so that even if acceptance criteria are limited, cases extrapolate all flows, error paths, user types, devices, and data inputs that might reasonably affect quality or outcomes.
6. If relevant, suggest test cases for workflows across modules and for special user profiles or edge hardware/software settings.
7. Return only the JSON array, no extra commentary.
"""
    return prompt.strip()

def call_groq_for_tests(prompt: str, api_key: str, model: str = "llama-3.3-70b-versatile", timeout: int = 120):
    if not GROQ_CLIENT_AVAILABLE:
        raise RuntimeError("Groq client not available. Install: pip install groq")
    
    client = Groq(api_key=api_key)
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_tokens=32768,
        timeout=timeout,
    )
    text = response.choices[0].message.content.strip()

    # Sanitize the output to fix JSON issues before loading
    sanitized_text = sanitize_groq_output(text)

    try:
        return json.loads(sanitized_text)
    except Exception:
        first, last = sanitized_text.find("["), sanitized_text.rfind("]")
        if first != -1 and last != -1:
            try:
                return json.loads(sanitized_text[first:last+1])
            except Exception as e:
                raise ValueError(f"Groq output not valid JSON after sanitizing:\n{sanitized_text[:1000]}\nError: {str(e)}")
        raise ValueError(f"Groq output not valid JSON:\n{sanitized_text[:1000]}")

def story_to_testcases(title: str, description_html: str, acceptance_html: str, groq_api_key: str, model: str):
    desc_text = strip_html_to_text(description_html)
    ac_text = strip_html_to_text(acceptance_html)

    prompt = make_groq_test_prompt(title, desc_text, ac_text)
    ai_results = call_groq_for_tests(prompt, api_key=groq_api_key, model=model)

    rows = []
    for idx, item in enumerate(ai_results, start=1):
        tc_id = item.get("test_case_id") or f"TC_{idx}"
        priority = item.get("priority") or "Medium"
        tc_title = item.get("test_case") or ""
        pre = item.get("precondition") or ""
        steps = item.get("test_steps") or []
        if isinstance(steps, str):
            steps = [s.strip() for s in steps.splitlines() if s.strip()]
        expected = item.get("expected_result") or ""

        rows.append({
            "Test Case ID": tc_id,
            "Priority": priority,
            "Test Case": tc_title,
            "Pre-condition / Test Data": pre,
            "Test Steps": "<br>".join([f"{i+1}. {s}" for i, s in enumerate(steps)]),
            "Expected Result": expected
        })

    return pd.DataFrame(rows, columns=[
        "Test Case ID", "Priority", "Test Case",
        "Pre-condition / Test Data", "Test Steps", "Expected Result"
    ])

def replace_br_with_newline_in_memory(wb):
    """
    Replace <br> tags with newlines in all cells of the workbook
    """
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and "<br>" in cell.value:
                    cell.value = cell.value.replace("<br>", "\n")
    
    return wb

def generate_test_cases_excel(story_ids: List[str], org: str, project: str, pat: str, groq_key: str, model: str) -> BytesIO:
    """
    Generate test cases for multiple stories and return Excel file as BytesIO
    """
    if pd is None or openpyxl is None:
        raise RuntimeError("pandas and openpyxl are required. Install: pip install pandas openpyxl")
    
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for sid in story_ids:
            try:
                desc, ac = azure_fetch_work_item(org, project, sid, pat)
                
                # Fetch title for better sheet naming
                org_clean = clean_org_for_api(org)
                base_url = f"https://dev.azure.com/{org_clean}"
                url = f"{base_url}/{project}/_apis/wit/workitems/{sid}?api-version=7.1"
                resp = requests.get(url, auth=("", pat))
                resp.raise_for_status()
                data = resp.json()
                title = data.get("fields", {}).get("System.Title", f"WorkItem_{sid}")
                
                st.info(f"Generating test cases for {sid} - {title}")
                
                df = story_to_testcases(title, desc, ac, groq_key, model)
            
                safe_sheet = f"{sid}"[:31]  # Excel sheet name limit
                df.to_excel(writer, sheet_name=safe_sheet, index=False)
                time.sleep(0.5)  # Rate limiting
                
            except Exception as e:
                st.warning(f"Failed to generate test cases for {sid}: {e}")
                # Create fallback sheet
                df = pd.DataFrame([{
                    "Test Case ID": "NA",
                    "Priority": "NA",
                    "Test Case": f"Generation failed for {sid}",
                    "Pre-condition / Test Data": "",
                    "Test Steps": str(e)[:500],
                    "Expected Result": ""
                }])
                safe_sheet = f"{sid}"[:31]
                df.to_excel(writer, sheet_name=safe_sheet, index=False)
    
    # Load the workbook and replace <br> with newlines
    output.seek(0)
    wb = openpyxl.load_workbook(output)
    wb = replace_br_with_newline_in_memory(wb)
    
    # Save back to BytesIO
    final_output = BytesIO()
    wb.save(final_output)
    final_output.seek(0)
    
    return final_output


def generate_test_cases_excel_for_jira(story_ids: List[str], base_url: str, pat: str, groq_key: str, model: str, email: str = None) -> BytesIO:
    """
    Generate test cases for Jira issues and return Excel file as BytesIO
    Uses jira_fetch_issue to get summary/description for each issue key.
    """
    if pd is None or openpyxl is None:
        raise RuntimeError("pandas and openpyxl are required. Install: pip install pandas openpyxl")

    output = BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for sid in story_ids:
            try:
                if jira_fetch_issue is None:
                    raise RuntimeError("Jira helpers not available. Check jira_integration.py import.")
                title, desc = jira_fetch_issue(base_url, sid, pat, email)
                ac = ""
                st.info(f"Generating test cases for {sid} - {title}")
                df = story_to_testcases(title, desc, ac, groq_key, model)
                # Smart Focus analysis
                smart_focus = generate_jira_pr_test_focus(
                    base_url=base_url,
                    issue_key=sid,   # here sid is defined by the loop
                    pat=pat,
                    story_text=desc,
                    email=email
                )

                st.markdown("### 🧠 Smart Testing Focus")
                st.markdown(smart_focus)
                safe_sheet = f"{sid}"[:31]
                df.to_excel(writer, sheet_name=safe_sheet, index=False)
                time.sleep(0.5)
            except Exception as e:
                st.warning(f"Failed to generate test cases for {sid}: {e}")
                df = pd.DataFrame([{
                    "Test Case ID": "NA",
                    "Priority": "NA",
                    "Test Case": f"Generation failed for {sid}",
                    "Pre-condition / Test Data": "",
                    "Test Steps": str(e)[:500],
                    "Expected Result": ""
                }])
                safe_sheet = f"{sid}"[:31]
                df.to_excel(writer, sheet_name=safe_sheet, index=False)

    output.seek(0)
    wb = openpyxl.load_workbook(output)
    wb = replace_br_with_newline_in_memory(wb)
    final_output = BytesIO()
    wb.save(final_output)
    final_output.seek(0)
    return final_output

def generate_jira_pr_test_focus(base_url: str, issue_key: str, pat: str, story_text: str, email: str = None) -> str:
    """
    Analyze linked PRs and generate Smart Testing Focus report using LLM.
    """
    try:
        prs = jira_fetch_linked_prs(base_url, issue_key, pat, email)
        
        if not prs:
            # Generate analysis based on story text alone when no PRs are linked
            prompt = f"""
You are an expert SDET and test architect.
Based on the following STORY CONTEXT (no code changes available),
analyze what areas would need testing and potential risk zones.

STORY CONTEXT:
{story_text[:1500]}

Generate your analysis in this format:
### 🧠 Smart Testing Focus (Story-Based Analysis)

**Note:** No linked PRs found. Analysis based on story requirements only.

**Key Test Areas**
- <area 1>
- <area 2>

**Potential Risk Zones**
- <risk 1>
- <risk 2>

**Suggested Test Scenarios**
- <scenario 1>
- <scenario 2>

**Regression Considerations**
- <consideration 1>
- <consideration 2>
"""
            llm = get_llm_instance()
            raw = llm.invoke(prompt)
            return extract_llm_text(raw)
        
        # Build PR summary
        changes_text = "\n".join([
            f"- **{pr['status']}**: {pr['title']}\n  URL: {pr['url']}\n  Author: {pr['author']}" 
            for pr in prs
        ])

        prompt = f"""
You are an expert SDET and test architect.
Based on the following STORY CONTEXT and LINKED PULL REQUESTS,
analyze where the system is most likely to break and what areas need deep testing.

STORY CONTEXT:
{story_text[:1500]}

LINKED PULL REQUESTS:
{changes_text}

Generate your analysis in this format:
### 🧠 Smart Testing Focus

**High Risk Areas**
- <short description>

**Potential Break Points**
- <short description>

**Suggested Smart Test Cases**
- <test case 1>
- <test case 2>

**Regression Scenarios**
- <scenario 1>
"""
        llm = get_llm_instance()
        raw = llm.invoke(prompt)
        return extract_llm_text(raw)
        
    except Exception as e:
        return f"""
### ⚠️ PR Analysis Error

Could not complete PR analysis: {str(e)[:200]}

**Possible reasons:**
- Jira dev-status API may not be enabled for your instance
- GitHub/Bitbucket integration may not be configured
- Required permissions may be missing

**Recommendation:** Focus on story-based testing approach without PR details.
"""
# -------------------------
# App UI & flow
# -------------------------

st.set_page_config(page_title="Story Summarizer (Multi) + QnA + Test Cases", layout="wide")
load_env_to_session()

st.title("Jira Story Summarizer — Multi-issue / JQL / Multi-PDF + Q&A + Test Cases")

st.sidebar.header("Configuration / Credentials")
# Accept either full browse URL or base URL
jira_base_input = st.sidebar.text_input("Jira Base URL (e.g. https://yourorg.atlassian.net)", key="jira_base_url", help="You may paste the full browse URL or base host")
jira_pat_input = st.sidebar.text_input("Jira API Token (JIRA_PAT)", key="jira_pat", type="password")
jira_email_input = st.sidebar.text_input("Optional Jira Email (for Basic auth)", key="jira_email", help="If your Jira requires email+token, provide the email here.")
st.sidebar.caption("Provide JIRA_PAT (or set JIRA_PAT in .env). If required, set JIRA_EMAIL for Basic auth.")
st.sidebar.write("GROQ_API_KEY present: " + ("Yes" if st.session_state.get("groq_key") else "No"))

# Create main tabs
main_tab1, main_tab2 = st.tabs(["📋 Story Summarizer & Q&A", "🧪 Test Case Generator"])

# -------------------------
# TAB 1: Story Summarizer & Q&A (Original functionality)
# -------------------------
with main_tab1:
    mode = st.radio("Input mode", ["Story IDs (comma-separated)", "WIQL Query", "Upload files (PDF/DOCX/TXT)"])

    col1, col2 = st.columns([3,1])
    with col1:
        if mode == "Story IDs (comma-separated)":
            story_ids = st.text_input("Enter story/work item IDs (comma-separated)", placeholder="e.g., 123, 456, 789")
        elif mode == "WIQL Query":
            wiql_query = st.text_area("Enter WIQL query", placeholder="SELECT [System.Id] FROM WorkItems WHERE ...")
        else:
            uploaded_files = st.file_uploader("Upload files (PDF, DOCX, TXT). You can upload multiple.", accept_multiple_files=True, type=["pdf","docx","txt","md"])

    with col2:
        st.markdown("### Options")
        append_mode = st.checkbox("Append to existing embedded corpus (don't overwrite)", value=False)
        generate_diagram = st.checkbox("Generate Flow Diagram", value=True, help="Create visual workflow diagram")
        process_btn = st.button("Fetch / Embed & Summarize")

    status = st.empty()
    embed_model_name = "sentence-transformers/all-MiniLM-L6-v2"

    if process_btn:
        try:
            status.info("Collecting inputs...")
            title_text_pairs = []
            if mode == "Story IDs (comma-separated)":
                if not story_ids or not story_ids.strip():
                    st.warning("Please enter at least one story ID.")
                else:
                    ids = [s.strip() for s in re.split(r"[,\s]+", story_ids.strip()) if s.strip()]
                    status.info(f"Fetching {len(ids)} issues from Jira...")
                    for sid in ids:
                        try:
                            if jira_fetch_issue is None:
                                raise RuntimeError("Jira helpers not available. Check jira_integration.py import.")
                            summary, desc = jira_fetch_issue(jira_base_input or st.session_state.get("jira_base_url"), sid, jira_pat_input or st.session_state.get("jira_pat"), jira_email_input or st.session_state.get("jira_email"))
                            ac = ""
                            comments = []
                            try:
                                comments = jira_fetch_comments(jira_base_input or st.session_state.get("jira_base_url"), sid, jira_pat_input or st.session_state.get("jira_pat"), jira_email_input or st.session_state.get("jira_email"))
                            except Exception:
                                comments = []

                            if comments:
                                formatted_comments = "\n".join([f"- [{c.get('author')}] ({c.get('created')}): {c.get('text')}" for c in comments])
                            else:
                                formatted_comments = "No comments found."

                            combined = "\n\n".join([
                                ("Summary:\n" + summary) if summary else "",
                                ("Description:\n" + desc) if desc else "",
                                ("Acceptance Criteria:\n" + ac) if ac else "",
                                ("Comments:\n" + formatted_comments) if formatted_comments else ""
                            ])
                            display_title = build_jira_issue_key(sid) if build_jira_issue_key else sid
                            title_text_pairs.append((f"{display_title}", combined.strip()))

                        except Exception as e:
                            st.warning(f"Failed to fetch {sid}: {e}")
            elif mode == "JQL Query" or mode == "WIQL Query":
                # Support JQL queries for Jira. If user left old WIQL label, still attempt to run as JQL.
                query = wiql_query if 'wiql_query' in locals() else ""
                if not query or not query.strip():
                    st.warning("Please provide a JQL query.")
                else:
                    status.info("Running JQL query...")
                    try:
                        ids = jira_run_jql(jira_base_input or st.session_state.get("jira_base_url"), query, jira_pat_input or st.session_state.get("jira_pat"), jira_email_input or st.session_state.get("jira_email"))
                        st.info(f"JQL returned {len(ids)} issues.")
                        for sid in ids:
                            try:
                                summary, desc = jira_fetch_issue(jira_base_input or st.session_state.get("jira_base_url"), sid, jira_pat_input or st.session_state.get("jira_pat"), jira_email_input or st.session_state.get("jira_email"))
                                ac = ""
                                combined = "\n\n".join([("Summary:\n" + summary) if summary else "", ("Description:\n" + desc) if desc else "", ("Acceptance Criteria:\n" + ac) if ac else ""])
                                title_text_pairs.append((f"{sid}", combined.strip()))
                            except Exception as e:
                                st.warning(f"Failed to fetch {sid}: {e}")
                    except Exception as e:
                        st.warning(f"JQL query failed: {e}")
            else:
                if not uploaded_files:
                    st.warning("Please upload at least one file.")
                else:
                    for f in uploaded_files:
                        try:
                            text = extract_text_from_file(f)
                            title_text_pairs.append((f.name, text))
                        except Exception as e:
                            st.warning(f"Failed to read {f.name}: {e}")

            if not title_text_pairs:
                status.error("No content collected to embed. Fix errors above and try again.")
            else:
                st.session_state["story_data"] = title_text_pairs
                
                docs = build_docs_from_texts(title_text_pairs)
                if not docs:
                    status.error("No valid text extracted from inputs.")
                else:
                    status.info("Splitting & embedding documents into FAISS...")
                    split_docs = []
                    for d in docs:
                        chunks = strict_split(d.page_content, chunk_size=1000, chunk_overlap=200)
                        for i, c in enumerate(chunks):
                            split_docs.append(Document(page_content=c, metadata={"source": d.metadata.get("source"), "chunk": i}))
                    embeddings = HuggingFaceEmbeddings(model_name=embed_model_name)
                    new_vs = FAISS.from_documents(split_docs, embeddings)
                    
                    if append_mode and st.session_state.get("vectorstore") is not None:
                        st.session_state["vectorstore"] = new_vs
                    else:
                        st.session_state["vectorstore"] = new_vs
                        
                    st.session_state["docs_sources"] = [t for t, _ in title_text_pairs]
                    
                    try:
                        llm = get_llm_instance()
                        if st.session_state.get("retrieval_chain") is not None:
                            memory = st.session_state["retrieval_chain"].memory
                        else:
                            memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
                        
                        retriever = st.session_state["vectorstore"].as_retriever(search_kwargs={"k": 4})
                        conv_chain = ConversationalRetrievalChain.from_llm(llm=llm, retriever=retriever, memory=memory)
                        st.session_state["retrieval_chain"] = conv_chain
                        status.success("Embeddings created and LLM chain initialized.")
                    except Exception as e:
                        st.session_state["retrieval_chain"] = None
                        status.warning(f"Embeddings created but LLM not initialized: {e}")
                    
                    if generate_diagram and len(title_text_pairs) > 0:
                        status.info("Generating flow diagram...")
                        try:
                            flow_diagram = generate_flow_diagram(title_text_pairs)
                            st.session_state["flow_diagram"] = flow_diagram
                        except Exception as e:
                            st.warning(f"Could not generate flow diagram: {e}")
                    
                    status.info("Generating summary...")
                    combined_text = "\n\n".join([f"{title}\n\n{text[:5000]}" for title, text in title_text_pairs])
                    summary_prompt = f"""
You are an expert Agile analyst and QA architect.
Summarize the following Jira issues clearly in Markdown format.

Make sure to:
1. Extract main objectives from Description and Acceptance Criteria.
2. Analyze developer and QA comments for decisions, deferrals, or concerns.
3. Highlight any risk areas, incomplete work, or postponed validations.
4. Give a short "Discussion Summary" section for each story.

Each story summary should follow this format:

### 🧩 {{Story Title}}

**Overview**
- One or two lines explaining what this story is about.

**Layman Summary**
- Summarize the whole story in layman terms + bullet points

**P0 Scenarios**
- Tell the most critical and P0 scenarios which are mandatory

**Acceptance & Completion Criteria**
- Summarize the acceptance or done conditions.

**Comment Insights**
- Mention relevant comments or discussion takeaways.
- Prioritize developer or QA comments (⭐3 or ⭐2).

---

CONTENT START
{combined_text}
CONTENT END
"""

                    try:
                        llm_temp = get_llm_instance()
                        raw_out = llm_temp.invoke(summary_prompt)
                        layman_summary = extract_llm_text(raw_out)
                        st.session_state["layman_summary"] = layman_summary
                    except Exception as e:
                        st.session_state["layman_summary"] = f"(Summary not generated: {e})"
                        
                    status.success("Done!")
        except Exception as e:
            status.error(f"Failed: {e}")

    # Flow Diagram Section
    if st.session_state.get("flow_diagram"):
        st.markdown("---")
        st.header("📊 Visual Flow Diagram")
        
        if st.button("🔄 Regenerate", key="regen_btn"):
            if st.session_state.get("story_data"):
                try:
                    flow_diagram = generate_flow_diagram(st.session_state["story_data"])
                    st.session_state["flow_diagram"] = flow_diagram
                    st.rerun()
                except Exception as e:
                    st.error(f"Failed: {e}")
        
        diagram_code = st.session_state["flow_diagram"]
        
        # Clean diagram code
        if "content='" in diagram_code:
            match = re.search(r'(graph\s+\w+.*?)(?=\'|$)', diagram_code, re.DOTALL)
            if match:
                diagram_code = match.group(1)
        
        diagram_code = diagram_code.replace('\\n', '\n')
        diagram_code = re.sub(r'```.*?```', '', diagram_code, flags=re.DOTALL)
        diagram_code = diagram_code.strip()
        
        # Create tabs
        tab1, tab2 = st.tabs(["🎯 Visual Diagram", "📝 Code"])
        
        with tab1:
            # Render using Mermaid.js
            mermaid_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
                <script>
                    mermaid.initialize({{ startOnLoad: true, theme: 'default' }});
                </script>
                <style>
                    body {{ 
                        font-family: Arial, sans-serif;
                        margin: 20px;
                        background: white;
                    }}
                    .mermaid {{ text-align: center; }}
                </style>
            </head>
            <body>
                <div class="mermaid">
    {diagram_code}
                </div>
            </body>
            </html>
            """
            
            import streamlit.components.v1 as components
            components.html(mermaid_html, height=500, scrolling=True)
        
        with tab2:
            st.code(diagram_code, language="mermaid")
            st.download_button(
                label="💾 Download Mermaid Code",
                data=diagram_code,
                file_name="flow_diagram.mmd",
                mime="text/plain",
                key="download_btn"
            )

    # Layman Summary
    st.markdown("---")
    st.header("📄 Summary")

    if st.session_state.get("layman_summary"):
        summary = st.session_state["layman_summary"]
        if isinstance(summary, str):
            summary = summary.replace("\\n", "\n")
        st.markdown(summary, unsafe_allow_html=True)
    else:
        st.info("No summary yet.")

    # Q&A Section
    st.markdown("---")
    st.header("💬 Ask Questions")

    if st.session_state.get("vectorstore") is None:
        st.info("Embed content first.")
    elif st.session_state.get("retrieval_chain") is None:
        st.warning("LLM not initialized.")
    else:
        with st.form(key="qa_form", clear_on_submit=True):
            user_q = st.text_input("Ask a question:", key="q_input")
            submit_q = st.form_submit_button("Ask")

        if submit_q and user_q.strip():
            try:
                chain = st.session_state["retrieval_chain"]
                result = chain.invoke({"question": user_q})
                answer = extract_llm_text(result)
                st.session_state.qa_history.insert(0, (user_q, answer))
                st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")

        if st.session_state.qa_history:
            st.markdown(f"### 💬 History ({len(st.session_state.qa_history)} questions)")
            for idx, (q, a) in enumerate(st.session_state.qa_history):
                with st.expander(f"Q{len(st.session_state.qa_history) - idx}: {q}", expanded=(idx == 0)):
                    st.markdown(f"**Answer:**\n\n{a}")

    # SMART TESTING FOCUS / RISK ANALYSIS SECTION
    st.markdown("---")
    st.header("🧪 Smart Testing Focus & Risk Zones")

    if st.session_state.get("story_data") and st.session_state.get("vectorstore"):
        selected_story = st.selectbox(
            "Select a story to analyze for linked PRs:",
            options=[t for t, _ in st.session_state["story_data"]],
            help="Choose a story whose PR impact you want to analyze"
        )
        
        if st.button("🔍 Analyze Linked PRs & Suggest Test Focus"):
            story_text = next((text for title, text in st.session_state["story_data"] if title == selected_story), "")
            # Extract the issue key from selected_story (handles formats like "BAW-123" or just "BAW-123")
            issue_key = selected_story.split(":")[0].strip() if ":" in selected_story else selected_story.strip()
            
            with st.spinner("Analyzing PRs and generating smart testing insights..."):
                pr_analysis = generate_jira_pr_test_focus(
                    base_url=st.session_state["jira_base_url"],
                    issue_key=issue_key,  # ✅ Now using the correct variable
                    pat=st.session_state["jira_pat"],
                    story_text=story_text,  # ✅ Using story_text instead of undefined 'desc'
                    email=st.session_state["jira_email"]
                )
                st.session_state["pr_analysis"] = pr_analysis
                st.success("✅ Analysis complete!")

    if st.session_state.get("pr_analysis"):
        st.markdown(st.session_state["pr_analysis"], unsafe_allow_html=True)
    else:
        st.info("Select a story and click 'Analyze Linked PRs & Suggest Test Focus' to generate insights.")

    # Debug
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("📄 Sources", key="src_btn"):
            st.json(st.session_state.get("docs_sources", []))
    with col2:
        if st.button("🎨 Diagram", key="diag_btn"):
            if st.session_state.get("story_data"):
                flow_diagram = generate_flow_diagram(st.session_state["story_data"])
                st.session_state["flow_diagram"] = flow_diagram
                st.rerun()
    with col3:
        if st.button("🗑️ Reset", key="reset_btn"):
            for key in ["vectorstore", "retrieval_chain", "docs_sources", "layman_summary", "chat_history", "qa_history", "flow_diagram", "story_data", "pr_analysis"]:
                st.session_state.pop(key, None)
            st.rerun()

# -------------------------
# TAB 2: Test Case Generator (New functionality)
# -------------------------
with main_tab2:
    st.header("🧪 AI-Powered Test Case Generator")
    st.markdown("""
    Generate comprehensive test cases for Jira issues using AI.
    The system will create 30-100 test cases covering all scenarios including edge cases, negative testing, and more.
    """)
    
    # Check dependencies
    if not GROQ_CLIENT_AVAILABLE:
        st.error("⚠️ Groq client not installed. Run: `pip install groq`")
    if pd is None or openpyxl is None:
        st.error("⚠️ Excel libraries not installed. Run: `pip install pandas openpyxl`")
    
    test_mode = st.radio("Select input mode for test case generation:", 
                         ["Issue Keys (comma-separated)", "JQL Query"], 
                         key="test_mode")
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        if test_mode == "Issue Keys (comma-separated)":
            test_story_ids = st.text_input(
                "Enter issue keys or numeric IDs (comma-separated)", 
                placeholder="e.g., BAW-123, 456, BAW-789",
                key="test_story_ids"
            )
        else:
            test_wiql_query = st.text_area(
                "Enter JQL query", 
                placeholder="project = BAW AND status = \"Ready for QA\" ORDER BY priority DESC",
                key="test_wiql"
            )
    
    with col2:
        st.markdown("### Settings")
        test_model = st.selectbox(
            "AI Model",
            ["llama-3.3-70b-versatile", "openai/gpt-oss-20b", "groq/compound"],
            key="test_model"
        )
        generate_test_btn = st.button("🚀 Generate Test Cases", type="primary")
    
    test_status = st.empty()
    
    if generate_test_btn:
        if not st.session_state.get("groq_key"):
            test_status.error("❌ GROQ_API_KEY not found. Please set it in your .env file or sidebar.")
        elif not GROQ_CLIENT_AVAILABLE or pd is None or openpyxl is None:
            test_status.error("❌ Required libraries not installed. Check error messages above.")
        else:
            try:
                test_status.info("🔍 Collecting story IDs...")
                
                # Collect story IDs
                story_ids_list = []
                if test_mode == "Issue Keys (comma-separated)":
                    if not test_story_ids or not test_story_ids.strip():
                        test_status.warning("⚠️ Please enter at least one story ID.")
                    else:
                        story_ids_list = [s.strip() for s in re.split(r"[,\s]+", test_story_ids.strip()) if s.strip()]
                else:
                    if not test_wiql_query or not test_wiql_query.strip():
                        test_status.warning("⚠️ Please provide a JQL query.")
                    else:
                        test_status.info("🔍 Running JQL query...")
                        try:
                            story_ids_list = jira_run_jql(
                                jira_base_input or st.session_state.get("jira_base_url"),
                                test_wiql_query,
                                jira_pat_input or st.session_state.get("jira_pat"),
                                jira_email_input or st.session_state.get("jira_email")
                            )
                            test_status.info(f"✅ Found {len(story_ids_list)} issues")
                        except Exception as e:
                            test_status.error(f"JQL query failed: {e}")
                
                if not story_ids_list:
                    test_status.error("❌ No stories found to process.")
                else:
                    test_status.info(f"⚙️ Generating test cases for {len(story_ids_list)} stories...")
                    
                    # Generate test cases
                    excel_file = generate_test_cases_excel_for_jira(
                        story_ids_list,
                        jira_base_input or st.session_state.get("jira_base_url"),
                        jira_pat_input or st.session_state.get("jira_pat"),
                        st.session_state["groq_key"],
                        test_model,
                        jira_email_input or st.session_state.get("jira_email")
                    )
                    
                    st.session_state["test_cases_excel"] = excel_file
                    test_status.success(f"✅ Test cases generated successfully for {len(story_ids_list)} stories!")
                    
            except Exception as e:
                test_status.error(f"❌ Error: {str(e)}")
                import traceback
                st.error(traceback.format_exc())
    
    # Download section
    if st.session_state.get("test_cases_excel"):
        st.markdown("---")
        st.subheader("📥 Download Test Cases")
        
        col1, col2, col3 = st.columns([2, 1, 1])
        
        with col1:
            st.download_button(
                label="💾 Download Excel File",
                data=st.session_state["test_cases_excel"],
                file_name=f"test_cases_{time.strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_test_cases"
            )
        
        with col2:
            if st.button("🗑️ Clear Results", key="clear_test_results"):
                st.session_state["test_cases_excel"] = None
                st.rerun()
        
        st.info("ℹ️ The Excel file contains formatted test cases with proper line breaks. Each story has its own sheet.")
    
    # Instructions
    st.markdown("---")
    with st.expander("ℹ️ How to use Test Case Generator"):
                st.markdown("""
                ### Instructions
        
                1. **Configure Jira credentials** in the sidebar (Base URL, JIRA_PAT, optional JIRA_EMAIL)
                2. **Ensure GROQ_API_KEY** is set in your environment or .env file
                3. **Choose input mode:**
                     - **Issue Keys**: Enter comma-separated issue keys or numeric IDs (e.g., BAW-123, 456)
                     - **JQL Query**: Write a custom JQL query to fetch multiple issues
                4. **Select AI model** (llama-3.3-70b-versatile recommended)
                5. **Click "Generate Test Cases"** and wait for processing
                6. **Download the Excel file** with comprehensive test cases

                ### What you get
                - 30-100 test cases per issue covering:
                    - ✅ Positive scenarios
                    - ❌ Negative scenarios  
                    - 🔄 Edge cases and boundary values
                    - 🔗 Integration scenarios
                    - 🔐 Security test cases
                    - ♿ Accessibility checks
                    - 📱 Compatibility testing
        
                ### Excel Format
                Each sheet contains:
                - Test Case ID
                - Priority (High/Medium/Low)
                - Test Case title
                - Pre-conditions/Test Data
                - Test Steps (formatted with line breaks)
                - Expected Results
                """)